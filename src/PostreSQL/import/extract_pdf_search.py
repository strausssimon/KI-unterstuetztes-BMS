import tkinter as tk
from tkinter import filedialog, messagebox
import os
from pathlib import Path
from sentence_transformers import SentenceTransformer, util
import PyPDF2
import docx
import re
from fuzzywuzzy import fuzz
from docx.shared import RGBColor
import openai

# OpenAI API-Schlüssel setzen
openai.api_key = "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"

# Initialisiere das NLP-Modell
model = SentenceTransformer('all-MiniLM-L6-v2')

def select_save_path():
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Documents", "*.docx")],
        title="Speicherort wählen"
    )
    if file_path:
        save_path_var.set(file_path)

def select_directory():
    directory = filedialog.askdirectory()
    if directory:
        directory_var.set(directory)

def extract_text_from_file(file_path):
    if file_path.suffix == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif file_path.suffix == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_path.suffix == '.docx':
        return extract_text_from_docx(file_path)
    return None

def extract_text_from_pdf(pdf_path):
    text = []
    with open(pdf_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page_num, page in enumerate(reader.pages):
            text.append(f"[Seite {page_num + 1}] \n" + page.extract_text())
    return "\n".join(text)

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)

def clean_text(text):
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\x00-\x7F]+', '', text)
    return text

def highlight_query_in_text(text, query):
    return re.sub(r'(?i)\b' + re.escape(query) + r'\w*\b', lambda m: f'!{m.group(0).upper()}!', text)

def semantic_search(text, query, model, top_k=10, threshold=0.4):
    query_embedding = model.encode(query, convert_to_tensor=True)
    text_embedding = model.encode(text, convert_to_tensor=True)
    similarities = util.pytorch_cos_sim(query_embedding, text_embedding).numpy().flatten()
    top_results = similarities.argsort()[-top_k:][::-1]
    matches = [{"quote": text[idx], "similarity": similarities[idx]} for idx in top_results if similarities[idx] > threshold]
    return matches

def enhanced_semantic_search(text, query, model, gpt_enabled=True, top_k=10, threshold=0.4):
    semantic_matches = semantic_search(text, query, model, top_k, threshold)
    gpt_results = None
    if gpt_enabled:
        gpt_analysis = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": f"Analysiere den folgenden Text basierend auf '{query}': {text[:2000]}"}]
        )
        gpt_results = gpt_analysis['choices'][0]['message']['content']
    return semantic_matches, gpt_results

def syntactic_search(text, query, threshold=70):
    sentences = re.split(r'(?<=[.!?])\s+', text)
    return [sentence for sentence in sentences if fuzz.partial_ratio(sentence.lower(), query.lower()) >= threshold]

def search_directory_with_gpt(directory, query, model, top_k=10, fuzzy_threshold=70, semantic_threshold=0.4, gpt_enabled=True):
    results, seen_documents = [], set()
    for root, _, files in os.walk(directory):
        for file in files:
            file_path = Path(root) / file
            if file_path.suffix in ['.txt', '.pdf', '.docx']:
                text = extract_text_from_file(file_path)
                if text:
                    text = clean_text(text)
                    syntactic_matches = syntactic_search(text, query, fuzzy_threshold)
                    semantic_matches, gpt_results = enhanced_semantic_search(text, query, model, gpt_enabled, top_k, semantic_threshold)
                    
                    for match in syntactic_matches:
                        highlighted_match = highlight_query_in_text(match, query)
                        results.append({"file": str(file_path), "quote": highlighted_match})
                    
                    for match in semantic_matches:
                        if str(file_path) not in seen_documents:
                            results.append({"file": str(file_path), "quote": match["quote"], "similarity": match["similarity"]})
                            seen_documents.add(str(file_path))
                    
                    if gpt_results:
                        results.append({"file": str(file_path), "quote": gpt_results})
    return results

def save_to_word(results, output_file):
    doc = docx.Document()
    doc.add_heading('Suchergebnisse', 0)
    for result in results:
        doc.add_heading(f"Datei: {result['file']}", level=1)
        p = doc.add_paragraph()
        words = result['quote'].split()
        for word in words:
            if word.startswith('!') and word.endswith('!'):
                run = p.add_run(word)
                run.font.color.rgb = RGBColor(255, 0, 0)
            else:
                p.add_run(word + " ")
        doc.add_paragraph("-" * 80)
    doc.save(output_file)
    messagebox.showinfo("Erfolg", f"Ergebnisse gespeichert in {output_file}")

def start_analysis():
    directory, query, gpt_enabled = directory_var.get(), query_var.get(), gpt_var.get()
    if not directory or not query:
        messagebox.showerror("Fehler", "Bitte geben Sie ein Verzeichnis und einen Suchbegriff ein.")
        return
    try:
        results = search_directory_with_gpt(directory, query, model, gpt_enabled=gpt_enabled)
        if results:
            save_to_word(results, save_path_var.get())
        else:
            messagebox.showinfo("Keine Treffer", "Keine passenden Ergebnisse gefunden.")
    except Exception as e:
        messagebox.showerror("Fehler", f"Ein Fehler ist aufgetreten: {e}")

# GUI-Elemente erstellen
root = tk.Tk()
root.title("Dateianalyse Tool")

tk.Label(root, text="Verzeichnis:").grid(row=0, column=0, padx=10, pady=10, sticky="w")
directory_var = tk.StringVar()
tk.Entry(root, textvariable=directory_var, width=50).grid(row=0, column=1, padx=10, pady=10)
tk.Button(root, text="Wählen", command=select_directory).grid(row=0, column=2, padx=10, pady=10)

tk.Label(root, text="Suchbegriff:").grid(row=1, column=0, padx=10, pady=10, sticky="w")
query_var = tk.StringVar()
tk.Entry(root, textvariable=query_var, width=50).grid(row=1, column=1, padx=10, pady=10)

gpt_var = tk.BooleanVar()
tk.Checkbutton(root, text="GPT-Analyse verwenden", variable=gpt_var).grid(row=2, column=1, pady=10, sticky="w")

tk.Label(root, text="Speicherort:").grid(row=3, column=0, padx=10, pady=10, sticky="w")
save_path_var = tk.StringVar()
tk.Entry(root, textvariable=save_path_var, width=50).grid(row=3, column=1, padx=10, pady=10)
tk.Button(root, text="Wählen", command=select_save_path).grid(row=3, column=2, padx=10, pady=10)

tk.Button(root, text="Analyse starten", command=start_analysis).grid(row=4, column=0, columnspan=3, pady=20)

root.mainloop()